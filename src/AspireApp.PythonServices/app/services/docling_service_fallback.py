import json
import os
from pathlib import Path
from typing import List, Dict, Any
from datetime import datetime

# Try to import docling, fall back to basic PDF processing if not available
try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.backend.pypdfium2_backend import PyPdfiumDocumentBackend
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False
    print("??  Docling not available, using fallback PDF processing")

# Fallback imports
try:
    import PyPDF2 as pypdf2
    PYPDF2_AVAILABLE = True
except ImportError:
    try:
        import pypdf2
        PYPDF2_AVAILABLE = True
    except ImportError:
        PYPDF2_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ..models.models import Document, ProcessedDocument, PageContent


class DoclingService:
    def __init__(self, data_path: str = "/app/data"):
        self.data_path = Path(data_path)
        self.processed_path = self.data_path / "processed"
        self.uploads_path = self.data_path / "uploads"
        
        # Ensure directories exist
        self.processed_path.mkdir(parents=True, exist_ok=True)
        (self.processed_path / "documents").mkdir(parents=True, exist_ok=True)
        
        # Initialize the document converter if available
        if DOCLING_AVAILABLE:
            try:
                self.converter = DocumentConverter()
                self.use_docling = True
                print("? Using Docling for document processing")
            except Exception as e:
                print(f"??  Docling initialization failed: {e}, using fallback")
                self.use_docling = False
        else:
            self.use_docling = False
            print("?? Using fallback document processing")
            print(f"   - PDF support: {'?' if PYPDF2_AVAILABLE else '?'}")
            print(f"   - DOCX support: {'?' if DOCX_AVAILABLE else '?'}")

    def process_document(self, document: Document) -> tuple[ProcessedDocument, List[PageContent]]:
        """Process a document using docling or fallback processors"""
        try:
            # Get the full file path
            file_path = self.uploads_path / document.file_path
            
            if not file_path.exists():
                raise FileNotFoundError(f"Document file not found: {file_path}")
            
            # Create document-specific directory
            doc_dir = self.processed_path / "documents" / str(document.id)
            doc_dir.mkdir(parents=True, exist_ok=True)
            
            # Create pages directory
            pages_dir = doc_dir / "pages"
            pages_dir.mkdir(exist_ok=True)
            
            # Process based on available libraries and file type
            if self.use_docling:
                return self._process_with_docling(document, file_path, doc_dir, pages_dir)
            else:
                return self._process_with_fallback(document, file_path, doc_dir, pages_dir)
                
        except Exception as e:
            raise RuntimeError(f"Failed to process document {document.id}: {str(e)}")

    def _process_with_docling(self, document: Document, file_path: Path, doc_dir: Path, pages_dir: Path) -> tuple[ProcessedDocument, List[PageContent]]:
        """Process document using full docling capabilities"""
        # Convert the document
        doc_result = self.converter.convert(str(file_path))
        docling_doc = doc_result.document
        
        # Save the full docling document
        document_json_path = doc_dir / "document.json"
        with open(document_json_path, 'w', encoding='utf-8') as f:
            json.dump(docling_doc.export_to_dict(), f, indent=2, ensure_ascii=False)
        
        # Extract pages and content
        pages = self._extract_pages_docling(docling_doc, pages_dir)
        
        # Create metadata
        metadata = {
            "processor": "docling",
            "processing_timestamp": datetime.now().isoformat(),
            "file_size": document.file_size,
            "mime_type": document.mime_type,
            "total_pages": len(pages)
        }
        
        # Save metadata
        metadata_path = doc_dir / "metadata.json"
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=2)
        
        # Create processed document record
        processed_doc = ProcessedDocument(
            document_id=document.id,
            docling_document_path=str(document_json_path),
            total_pages=len(pages),
            processing_date=datetime.now(),
            processing_metadata=metadata
        )
        
        return processed_doc, pages

    def _process_with_fallback(self, document: Document, file_path: Path, doc_dir: Path, pages_dir: Path) -> tuple[ProcessedDocument, List[PageContent]]:
        """Process document using fallback processors"""
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf' and PYPDF2_AVAILABLE:
            pages = self._extract_pages_pdf(file_path, pages_dir)
            processor = "pypdf2"
        elif file_extension in ['.docx', '.doc'] and DOCX_AVAILABLE:
            pages = self._extract_pages_docx(file_path, pages_dir)
            processor = "python-docx"
        else:
            # Basic text extraction for any file
            pages = self._extract_pages_text(file_path, pages_dir)
            processor = "text-fallback"
        
        # Create basic document structure
        document_data = {
            "filename": document.filename,
            "original_filename": document.original_filename,
            "file_path": str(file_path),
            "processor": processor,
            "pages": [{"page_number": p.page_number, "content": p.content} for p in pages]
        }
        
        # Save the document
        document_json_path = doc_dir / "document.json"
        with open(document_json_path, 'w', encoding='utf-8') as f:
            json.dump(document_data, f, indent=2, ensure_ascii=False)
        
        # Create metadata
        metadata = {
            "processor": processor,
            "processing_timestamp": datetime.now().isoformat(),
            "file_size": document.file_size,
            "mime_type": document.mime_type,
            "total_pages": len(pages),
            "note": "Processed with fallback processor due to missing dependencies",
            "available_processors": {
                "pypdf2": PYPDF2_AVAILABLE,
                "python_docx": DOCX_AVAILABLE,
                "docling": DOCLING_AVAILABLE
            }
        }
        
        # Save metadata
        metadata_path = doc_dir / "metadata.json"
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=2)
        
        # Create processed document record
        processed_doc = ProcessedDocument(
            document_id=document.id,
            docling_document_path=str(document_json_path),
            total_pages=len(pages),
            processing_date=datetime.now(),
            processing_metadata=metadata
        )
        
        return processed_doc, pages

    def _extract_pages_docling(self, docling_doc, pages_dir: Path) -> List[PageContent]:
        """Extract pages using docling"""
        pages = []
        
        # Get page content from docling document
        page_items = []
        for item in docling_doc.iterate_items():
            if hasattr(item, 'page') and item.page is not None:
                page_items.append((item.page, item))
        
        # Group by page number
        pages_dict = {}
        for page_num, item in page_items:
            if page_num not in pages_dict:
                pages_dict[page_num] = []
            pages_dict[page_num].append(item)
        
        # Process each page
        for page_num in sorted(pages_dict.keys()):
            page_content = ""
            page_metadata = {
                "page_number": page_num,
                "items_count": len(pages_dict[page_num])
            }
            
            # Extract text content from page items
            for item in pages_dict[page_num]:
                if hasattr(item, 'text') and item.text:
                    page_content += item.text + "\n"
            
            # Create page content object
            page = PageContent(
                page_number=page_num,
                content=page_content.strip(),
                metadata=page_metadata
            )
            
            # Save individual page
            page_file = pages_dir / f"page_{page_num:03d}.json"
            with open(page_file, 'w', encoding='utf-8') as f:
                json.dump({
                    "page_number": page_num,
                    "content": page_content.strip(),
                    "metadata": page_metadata
                }, f, indent=2, ensure_ascii=False)
            
            pages.append(page)
        
        return pages

    def _extract_pages_pdf(self, file_path: Path, pages_dir: Path) -> List[PageContent]:
        """Extract pages from PDF using PyPDF2"""
        pages = []
        
        try:
            with open(file_path, 'rb') as file:
                if hasattr(pypdf2, 'PdfReader'):
                    # New PyPDF2 API
                    pdf_reader = pypdf2.PdfReader(file)
                    pdf_pages = pdf_reader.pages
                else:
                    # Old PyPDF2 API
                    pdf_reader = pypdf2.PdfFileReader(file)
                    pdf_pages = [pdf_reader.getPage(i) for i in range(pdf_reader.numPages)]
                
                for page_num, page in enumerate(pdf_pages, 1):
                    try:
                        if hasattr(page, 'extract_text'):
                            content = page.extract_text()
                        else:
                            content = page.extractText()
                        
                        page_metadata = {
                            "page_number": page_num,
                            "processor": "pypdf2",
                            "api_version": "new" if hasattr(pypdf2, 'PdfReader') else "old"
                        }
                        
                        page_content = PageContent(
                            page_number=page_num,
                            content=content.strip(),
                            metadata=page_metadata
                        )
                        
                        # Save individual page
                        page_file = pages_dir / f"page_{page_num:03d}.json"
                        with open(page_file, 'w', encoding='utf-8') as f:
                            json.dump({
                                "page_number": page_num,
                                "content": content.strip(),
                                "metadata": page_metadata
                            }, f, indent=2, ensure_ascii=False)
                        
                        pages.append(page_content)
                        
                    except Exception as e:
                        print(f"Warning: Failed to extract page {page_num}: {e}")
                        continue
                        
        except Exception as e:
            raise RuntimeError(f"Failed to process PDF: {e}")
        
        return pages

    def _extract_pages_docx(self, file_path: Path, pages_dir: Path) -> List[PageContent]:
        """Extract content from DOCX using python-docx"""
        try:
            doc = DocxDocument(file_path)
            content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            page_metadata = {
                "page_number": 1,
                "processor": "python-docx",
                "paragraphs": len(doc.paragraphs)
            }
            
            page_content = PageContent(
                page_number=1,
                content=content,
                metadata=page_metadata
            )
            
            # Save page
            page_file = pages_dir / "page_001.json"
            with open(page_file, 'w', encoding='utf-8') as f:
                json.dump({
                    "page_number": 1,
                    "content": content,
                    "metadata": page_metadata
                }, f, indent=2, ensure_ascii=False)
            
            return [page_content]
            
        except Exception as e:
            raise RuntimeError(f"Failed to process DOCX: {e}")

    def _extract_pages_text(self, file_path: Path, pages_dir: Path) -> List[PageContent]:
        """Fallback text extraction for unsupported files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                # Final fallback - read as binary and decode with errors='ignore'
                with open(file_path, 'rb') as f:
                    content = f.read().decode('utf-8', errors='ignore')
            
            page_metadata = {
                "page_number": 1,
                "processor": "text-fallback",
                "file_size": file_path.stat().st_size,
                "encoding": "auto-detected"
            }
            
            page_content = PageContent(
                page_number=1,
                content=content[:10000] if len(content) > 10000 else content,  # Limit content size
                metadata=page_metadata
            )
            
            # Save page
            page_file = pages_dir / "page_001.json"
            with open(page_file, 'w', encoding='utf-8') as f:
                json.dump({
                    "page_number": 1,
                    "content": page_content.content,
                    "metadata": page_metadata
                }, f, indent=2, ensure_ascii=False)
            
            return [page_content]
            
        except Exception as e:
            raise RuntimeError(f"Failed to process text file: {e}")

    def get_document_path(self, doc_id: int) -> Path:
        """Get the path to a processed document directory"""
        return self.processed_path / "documents" / str(doc_id)

    def load_processed_document(self, doc_id: int) -> Dict[str, Any]:
        """Load a processed document"""
        doc_path = self.get_document_path(doc_id) / "document.json"
        
        if not doc_path.exists():
            raise FileNotFoundError(f"Processed document not found: {doc_path}")
        
        with open(doc_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    def load_page_content(self, doc_id: int, page_number: int) -> PageContent:
        """Load specific page content"""
        page_path = self.get_document_path(doc_id) / "pages" / f"page_{page_number:03d}.json"
        
        if not page_path.exists():
            raise FileNotFoundError(f"Page not found: {page_path}")
        
        with open(page_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return PageContent(**data)